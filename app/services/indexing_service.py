import threading
from datetime import timedelta

from flask import current_app

from ..extensions import db
from ..models import FileIndex, IndexJob, StoredFile, utcnow
from . import ai_service, file_service, search_service

MAX_TEXT_CHARS = 500_000
MAX_JOB_ATTEMPTS = 3


def extract_text(stored_file):
    """Extract searchable text from a file based on its extension."""
    ext = stored_file.extension
    path = file_service.file_path(stored_file)

    if ext == "pdf":
        return _extract_pdf(path)
    if ext == "docx":
        return _extract_docx(path)
    if ext in current_app.config["EDITABLE_EXTENSIONS"]:
        return file_service.read_text_content(stored_file, max_chars=MAX_TEXT_CHARS)
    if stored_file.is_image:
        return _extract_image_ocr(path)
    return ""


def _extract_pdf(path):
    import pdfplumber

    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    parts.append(" | ".join(cell or "" for cell in row))
            if sum(len(p) for p in parts) > MAX_TEXT_CHARS:
                break
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _extract_docx(path):
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)[:MAX_TEXT_CHARS]


_tesseract_missing_logged = False


def _tesseract_available():
    """Check once whether the tesseract binary is on PATH."""
    global _tesseract_missing_logged
    import shutil
    available = shutil.which("tesseract") is not None
    if not available and not _tesseract_missing_logged:
        _tesseract_missing_logged = True
        current_app.logger.info(
            "Tesseract not found on PATH — image OCR is disabled. "
            "Images will rely on AI captions (if an AI connection exists). "
            "Install tesseract to enable OCR.")
    return available


def _extract_image_ocr(path):
    from PIL import Image

    if not _tesseract_available():
        return ""
    import pytesseract

    langs = current_app.config.get("TESSERACT_LANGS", "eng")
    with Image.open(path) as img:
        text = pytesseract.image_to_string(img, lang=langs)
    return (text or "")[:MAX_TEXT_CHARS]


def index_file(file_id, app=None):
    """Index a single file: text extraction + AI caption for images."""
    app = app or current_app._get_current_object()
    with app.app_context():
        stored = db.session.get(StoredFile, file_id)
        if not stored:
            return
        index = stored.index or FileIndex(file_id=stored.id)
        if index.id is None:
            db.session.add(index)

        try:
            # OCR is best-effort for images: when tesseract is unavailable the
            # AI caption below still makes the image searchable.
            ocr_error = None
            try:
                text = extract_text(stored)
            except Exception as exc:  # noqa: BLE001
                if not stored.is_image:
                    raise
                ocr_error = exc
                text = ""
                app.logger.warning("OCR failed for image %s: %s", file_id, exc)

            caption = None
            ai_cfg = ai_service.config_for(stored.owner)
            # Synced drives can turn AI captions off (drive setting).
            captions_on = stored.drive.captions_enabled if stored.drive else True
            if stored.is_image and captions_on and ai_cfg["enabled"] and ai_cfg["base_url"]:
                caption = ai_service.caption_image(
                    file_service.file_path(stored), config=ai_cfg)

            if stored.is_image and ocr_error and not caption:
                raise ocr_error  # neither OCR nor AI caption produced anything

            index.extracted_text = text or None
            index.caption = caption or None
            # Content statistics (words/lines/chars) for the user to inspect
            stats_source = text or caption or ""
            index.char_count = len(stats_source) if stats_source else None
            index.word_count = len(stats_source.split()) if stats_source else None
            index.line_count = (stats_source.count("\n") + 1) if stats_source else None
            index.status = "ok"
            index.error = None
        except Exception as exc:  # noqa: BLE001 - record any extraction failure
            app.logger.exception("Indexing failed for file %s", file_id)
            index.status = "error"
            index.error = str(exc)[:1000]

        from ..models import utcnow
        index.indexed_at = utcnow()
        db.session.commit()
        # Keep the FTS5 search index in sync (no-op when FTS5 is unavailable).
        search_service.fts_upsert(stored.id)


# --------------------------------------------------------------------------
# Persistent index queue (index_jobs table)
#
# Uploads, edits, reindexes and folder syncs all enqueue jobs here instead of
# spawning ad-hoc threads. Jobs survive restarts: on startup, jobs left in
# "running" by a crash go back to "pending" and drainers resume the work.
# Failures are retried up to MAX_JOB_ATTEMPTS times, then marked "error"
# (the file's index badge shows the error too — nothing is silently lost).
# --------------------------------------------------------------------------

_claim_lock = threading.Lock()     # makes job claiming atomic (single process)
_workers_lock = threading.Lock()
_drainers_running = 0
_queue_paused = False


def enqueue_index(file_ids):
    """Persist pending jobs for the given files (deduped, skips missing or
    trashed files). Returns how many jobs were created."""
    created = 0
    for fid in file_ids:
        stored = db.session.get(StoredFile, fid)
        if not stored or stored.deleted_at is not None:
            continue
        exists = (IndexJob.query
                  .filter(IndexJob.file_id == fid,
                          IndexJob.status.in_(("pending", "running")))
                  .first())
        if exists:
            continue
        db.session.add(IndexJob(file_id=fid, status="pending"))
        created += 1
    db.session.commit()
    return created


def _claim_next_job():
    """Atomically claim the oldest pending job; returns the job id or None."""
    with _claim_lock:
        job = (IndexJob.query.filter_by(status="pending")
               .order_by(IndexJob.id).first())
        if not job:
            return None
        job.status = "running"
        job.attempts += 1
        db.session.commit()
        return job.id


def _drain(app):
    with app.app_context():
        while not _queue_paused:
            job_id = _claim_next_job()
            if job_id is None:
                return
            job = db.session.get(IndexJob, job_id)
            stored = db.session.get(StoredFile, job.file_id)
            if not stored or stored.deleted_at is not None:
                job.status = "done"  # file gone/trashed: nothing to do
                db.session.commit()
                continue
            # index_file records the outcome on FileIndex and never raises.
            index_file(stored.id, app)
            db.session.expire_all()  # index_file committed in its own context
            index = db.session.get(StoredFile, stored.id).index
            if index and index.status == "ok":
                job.status = "done"
                job.error = None
            else:
                job.error = (index.error if index else "indexing failed")
                # Retry while attempts remain; final failure stays visible.
                job.status = ("pending" if job.attempts < MAX_JOB_ATTEMPTS
                              else "error")
            db.session.commit()


def _drainer_entry(app):
    global _drainers_running
    try:
        _drain(app)
    except Exception:  # noqa: BLE001 - a drainer must never kill the app
        app.logger.exception("Index queue drainer crashed")
    finally:
        with _workers_lock:
            _drainers_running -= 1


def ensure_workers(n=1, app=None):
    """Spawn up to n drainer threads (clamped to 8) to work the queue."""
    global _drainers_running
    if _queue_paused:
        return
    app = app or current_app._get_current_object()
    want = min(max(1, int(n or 1)), 8)
    with _workers_lock:
        spawn = max(0, want - _drainers_running)
        _drainers_running += spawn
    for _ in range(spawn):
        threading.Thread(target=_drainer_entry, args=(app,), daemon=True).start()


def pause_queue():
    """Stop drainers from claiming new jobs. Pending jobs stay in the DB."""
    global _queue_paused
    _queue_paused = True


def resume_queue(app=None):
    """Clear the pause flag and spin drainers back up."""
    global _queue_paused
    _queue_paused = False
    app = app or current_app._get_current_object()
    ensure_workers(app.config.get("INDEX_WORKERS", 2), app)


def queue_paused():
    return _queue_paused


def get_queue_status(user):
    """Per-user queue view for the profile card: counts, progress, running
    file names and the global paused flag."""
    base = (IndexJob.query
            .join(StoredFile, StoredFile.id == IndexJob.file_id)
            .filter(StoredFile.user_id == user.id))
    counts = {"pending": 0, "running": 0, "done": 0, "error": 0}
    for status, n in (base.with_entities(IndexJob.status, db.func.count())
                      .group_by(IndexJob.status).all()):
        counts[status] = n
    running_files = [name for (name,) in
                     (base.filter(IndexJob.status == "running")
                      .with_entities(StoredFile.name).limit(5).all())]
    total = sum(counts.values())
    processed = counts["done"] + counts["error"]
    return {
        "paused": _queue_paused,
        "counts": counts,
        "running_files": running_files,
        "active": counts["pending"] + counts["running"],
        "processed": processed,
        "total": total,
        "percent": round(100 * processed / total) if total else 100,
    }


def recover_interrupted(app):
    """Startup recovery: requeue jobs a crash left "running", purge finished
    jobs older than 24h, and spin drainers if work remains."""
    with app.app_context():
        stale = IndexJob.query.filter_by(status="running")
        stale.update({"status": "pending"}, synchronize_session=False)
        cutoff = utcnow() - timedelta(hours=24)
        (IndexJob.query
         .filter(IndexJob.status.in_(("done", "error")),
                 IndexJob.updated_at < cutoff)
         .delete(synchronize_session=False))
        db.session.commit()
        pending = IndexJob.query.filter_by(status="pending").count()
    if pending and app.config.get("INDEX_ASYNC", True):
        # INDEX_ASYNC=False (tests): no real threads — pending jobs are
        # processed inline by the callers instead.
        ensure_workers(app.config.get("INDEX_WORKERS", 2), app)


def run_pending_inline():
    """Process every pending job in the current thread (tests /
    INDEX_ASYNC=False)."""
    while True:
        job_id = _claim_next_job()
        if job_id is None:
            return
        job = db.session.get(IndexJob, job_id)
        stored = db.session.get(StoredFile, job.file_id)
        if not stored or stored.deleted_at is not None:
            job.status = "done"
            db.session.commit()
            continue
        index_file(stored.id)
        db.session.expire_all()
        index = db.session.get(StoredFile, stored.id).index
        if index and index.status == "ok":
            job.status = "done"
            job.error = None
        else:
            job.error = (index.error if index else "indexing failed")
            job.status = ("pending" if job.attempts < MAX_JOB_ATTEMPTS
                          else "error")
            db.session.commit()
            if job.status == "pending":
                continue  # retry immediately in inline mode
        db.session.commit()
